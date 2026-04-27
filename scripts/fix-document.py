"""Fix all 14 observations in DOCUMENTO_PROYECTO_IA.docx"""
from docx import Document

doc = Document('DOCUMENTO_PROYECTO_IA.docx')

def replace_paragraph_text(doc, old_text, new_text):
    for para in doc.paragraphs:
        if old_text in para.text:
            if para.runs:
                para.runs[0].text = para.text.replace(old_text, new_text)
                for r in para.runs[1:]:
                    r.text = ''
            return True
    return False

def replace_paragraph_full(doc, search_fragment, new_text):
    """Replace entire paragraph text when fragment is found"""
    for para in doc.paragraphs:
        if search_fragment in para.text:
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ''
            return True
    return False

def replace_in_table_cell(doc, old_text, new_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text and para.runs:
                        para.runs[0].text = para.text.replace(old_text, new_text)
                        for r in para.runs[1:]:
                            r.text = ''
                        return True
    return False

def replace_exact_cell(doc, exact_text, new_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip() == exact_text and para.runs:
                        para.runs[0].text = new_text
                        for r in para.runs[1:]:
                            r.text = ''
                        return True
    return False

def replace_cell_in_row_with(doc, row_marker, old_cell, new_cell):
    """Find row containing row_marker, then replace old_cell with new_cell in that row"""
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            if any(row_marker in t for t in row_texts):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip() == old_cell and para.runs:
                            para.runs[0].text = new_cell
                            for r in para.runs[1:]:
                                r.text = ''
                            return True
    return False

changes = []

# === CORRECCIÓN 1: Problema Central ===
r = replace_paragraph_full(doc,
    'No existe una soluci',
    'Las instituciones de educaci\u00f3n t\u00e9cnica carecen de herramientas que integren, en una sola plataforma, tutor\u00eda adaptativa, dise\u00f1o instruccional asistido, evaluaci\u00f3n continua y monitoreo en tiempo real. Esta fragmentaci\u00f3n obliga a los docentes a gestionar m\u00faltiples sistemas desconectados, limitando su capacidad de ofrecer atenci\u00f3n personalizada a escala y dificultando la toma de decisiones pedag\u00f3gicas basadas en evidencia.'
)
changes.append(f'1. Problema Central: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 2: Obs calificación ===
r = replace_paragraph_text(doc,
    '(obs: cambiar al sistema actual)',
    ', alineada al sistema de evaluaci\u00f3n vigesimal del Ministerio de Educaci\u00f3n del Per\u00fa (nota aprobatoria: 13/20 = 65/100):'
)
changes.append(f'2. Obs calificaci\u00f3n: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 3: Comentario interno ===
r = replace_paragraph_full(doc,
    'actualmente tenemos a mas de 35',
    'Ratio instructor-estudiante desbalanceada: Un docente atiende entre 30 y 60 estudiantes simult\u00e1neamente, lo que impide ofrecer retroalimentaci\u00f3n individualizada. En el contexto de CETEMIN, con m\u00e1s de 35 estudiantes por sal\u00f3n, los instructores no logran brindar una ense\u00f1anza personalizada a pesar de sus esfuerzos.'
)
changes.append(f'3. Comentario interno: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 4: Tabla rúbrica - rangos ===
for old, new in [('90-100', '85-100'), ('70-89', '65-84'), ('50-69', '50-64')]:
    r = replace_exact_cell(doc, old, new)
    changes.append(f'4. Rango {old} -> {new}: {"OK" if r else "FAIL"}')

# Criterios de la tabla
for old, new in [
    ('An\u00e1lisis o aplicaci\u00f3n con 1-2 intentos', 'An\u00e1lisis o aplicaci\u00f3n demostrada con alta eficiencia (1-2 intentos)'),
    ('Comprensi\u00f3n demostrada en 3-4 intentos', 'Comprensi\u00f3n demostrada con eficiencia aceptable'),
    ('Comprensi\u00f3n parcial, m\u00faltiples intentos', 'Comprensi\u00f3n parcial, requiere refuerzo'),
    ('No completado o avance forzado', 'No logr\u00f3 demostrar comprensi\u00f3n o avanz\u00f3 por l\u00edmite de intentos'),
]:
    r = replace_in_table_cell(doc, old, new)
    changes.append(f'4. Criterio "{old[:30]}...": {"OK" if r else "FAIL"}')

# === CORRECCIÓN 5: KPI 70/100 ===
r = replace_in_table_cell(doc, '\u2265 70/100', '\u2265 65/100 (equivalente a 13/20)')
changes.append(f'5. KPI 70->65: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 6: Cronograma fechas ===
r = replace_exact_cell(doc, 'Semanas', 'Periodo')
changes.append(f'6. Header Semanas->Periodo: {"OK" if r else "FAIL"}')

for marker, old, new in [
    ('1. Fundaci', 'Sem 1-2', 'Sep 2025'),
    ('2. Chat IA', 'Sem 2-4', 'Oct - Nov 2025'),
    ('3. Inteligencia', 'Sem 4-6', 'Dic 2025'),
    ('4. Planificador', 'Sem 6-8', 'Ene - Feb 2026'),
    ('5. Pruebas', 'Sem 8-10', 'Mar - Abr 2026'),
]:
    r = replace_cell_in_row_with(doc, marker, old, new)
    changes.append(f'6. {old} -> {new}: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 7: Hitos fechas ===
r = replace_exact_cell(doc, 'Semana', 'Fecha')
changes.append(f'7. Header Semana->Fecha: {"OK" if r else "FAIL"}')

for marker, old, new in [
    ('H1', 'Sem 2', 'Sep 2025'),
    ('H2', 'Sem 4', 'Nov 2025'),
    ('H3', 'Sem 6', 'Dic 2025'),
    ('H4', 'Sem 8', 'Feb 2026'),
    ('H5', 'Sem 10', 'Abr 2026'),
]:
    r = replace_cell_in_row_with(doc, marker, old, new)
    changes.append(f'7. {marker} {old} -> {new}: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 7b: Gantt M1-M10 -> meses reales ===
month_map = {
    'M1': 'Sep', 'M2': 'Oct', 'M3': 'Nov', 'M4': 'Dic',
    'M5': 'Ene', 'M6': 'Feb', 'M7': 'Mar', 'M8': 'Abr',
    'M9': 'May', 'M10': 'Jun',
}
for old_m, new_m in month_map.items():
    r = replace_exact_cell(doc, old_m, new_m)
    if r:
        changes.append(f'7b. Gantt {old_m} -> {new_m}: OK')

# === CORRECCIÓN 8: Fecha portada ===
r = replace_paragraph_text(doc, '30 de marzo de 2026', 'Abril 2026')
changes.append(f'8. Fecha portada: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 10: Team name ===
r = replace_paragraph_text(doc, 'Sophia Next Development Team', 'Equipo de Desarrollo Sophia Next - CETEMIN')
changes.append(f'10. Team name: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 11: Caso de uso ===
r = replace_paragraph_full(doc,
    'Una universidad t\u00e9cnica necesita',
    'Contexto: CETEMIN necesita ense\u00f1ar la Ley 29783 de Seguridad y Salud en el Trabajo a estudiantes de la carrera t\u00e9cnica de Seguridad Industrial.'
)
changes.append(f'11. Caso de uso: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 12: Instructor -> docente CETEMIN ===
r = replace_paragraph_text(doc, 'El instructor ingresa al planificador', 'El docente de CETEMIN ingresa al planificador')
changes.append(f'12. Docente CETEMIN: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 13: Prácticas de ingeniería ===
practices = {
    'TypeScript estricto: Verificaci\u00f3n de tipos obligatoria antes de cada despliegue.':
        'Verificaci\u00f3n de c\u00f3digo estricta: Cada cambio pasa por validaci\u00f3n automatizada antes de publicarse, previniendo errores en producci\u00f3n.',
    'Server Components por defecto: Client Components solo cuando se requiere interactividad.':
        'Rendimiento optimizado: La interfaz prioriza velocidad de carga, procesando contenido en el servidor antes de enviarlo al navegador del estudiante.',
    'Prompt caching: Reutilizaci\u00f3n de prompts del sistema en ventanas de 5 minutos para reducir costos.':
        'Optimizaci\u00f3n de costos de IA: El sistema reutiliza instrucciones frecuentes del tutor, reduciendo consumo de recursos y costos operativos.',
    'Build de producci\u00f3n local: Obligatorio antes de push (npm run build).':
        'Control de calidad pre-publicaci\u00f3n: Cada actualizaci\u00f3n se compila y verifica antes de desplegarse a producci\u00f3n.',
}
for old, new in practices.items():
    r = replace_paragraph_text(doc, old, new)
    changes.append(f'13. Pr\u00e1ctica "{old[:40]}...": {"OK" if r else "FAIL"}')

# Streaming practice (may have different text)
r = replace_paragraph_full(doc,
    'Streaming de respuestas IA',
    'Respuestas en tiempo real: Las respuestas del tutor IA se muestran de forma progresiva (palabra por palabra), similar a ChatGPT, evitando tiempos de espera prolongados.'
)
changes.append(f'13. Streaming practice: {"OK" if r else "FAIL"}')

# === CORRECCIÓN 14: Architecture table simplify ===
arch = {
    'Interfaz de usuario con Server/Client Components': 'Aplicaci\u00f3n web responsiva con carga r\u00e1pida',
    'Chat Streaming, Planner, Dashboard, Admin, Auth': 'L\u00f3gica de negocio, chat, evaluaci\u00f3n y administraci\u00f3n',
    'Tutor\u00eda conversacional + Clasificaci\u00f3n de intenciones': 'Motor de tutor\u00eda conversacional y clasificaci\u00f3n de respuestas',
    'Abstracci\u00f3n de base de datos y migraciones': 'Gesti\u00f3n de estructura de datos',
    'Almacenamiento relacional en la nube': 'Almacenamiento de datos acad\u00e9micos, evidencias y calificaciones',
    'Despliegue, CDN global y serverless functions': 'Despliegue en la nube con disponibilidad global',
    'OAuth y gesti\u00f3n de sesiones JWT': 'Inicio de sesi\u00f3n seguro con Google',
}
for old, new in arch.items():
    r = replace_in_table_cell(doc, old, new)
    changes.append(f'14. Arch "{old[:35]}...": {"OK" if r else "FAIL"}')

doc.save('DOCUMENTO_PROYECTO_IA.docx')

print('\n=== RESULTADOS ===')
ok_count = sum(1 for c in changes if 'OK' in c)
fail_count = sum(1 for c in changes if 'FAIL' in c)
for c in changes:
    print(c)
print(f'\nTotal: {ok_count} OK, {fail_count} FAIL')
print('=== DOCUMENTO GUARDADO ===')
